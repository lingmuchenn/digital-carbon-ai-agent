import os
import sys
from pathlib import Path
from itertools import combinations
from typing import Any, Dict, List, Tuple
from datetime import datetime
import json
from difflib import SequenceMatcher
import re

import numpy as np
from docx import Document
from openai import OpenAI
from tqdm import tqdm

# 设置输出编码，避免 Windows 控制台默认 GBK 导致 emoji/特殊字符打印崩溃
if sys.platform == 'win32':
    try:
        sys.stdout.reconfigure(encoding='utf-8', errors='replace')
        sys.stderr.reconfigure(encoding='utf-8', errors='replace')
    except Exception:
        pass

def _safe_print(msg: str):
    """在控制台编码不支持时也不崩溃（打包/命令行场景常见）。"""
    try:
        print(msg)
    except UnicodeEncodeError:
        # 去掉可能导致崩溃的字符（如 emoji）
        print(msg.encode("utf-8", errors="replace").decode("utf-8", errors="replace"))

# 引入 ImageAnalyzer
try:
    from image_analysis import ImageAnalyzer
except Exception:
    try:
        # 兼容以包方式导入（import document_sort.find_similar_files）
        from .image_analysis import ImageAnalyzer  # type: ignore
    except Exception:
        ImageAnalyzer = None
        _safe_print("⚠️ 无法导入 image_analysis 模块，图片深度分析功能将不可用。")

try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

try:
    import pypdfium2 as pdfium
except ImportError:
    pdfium = None

try:
    from PIL import Image
except ImportError:
    Image = None

import base64
from io import BytesIO

# ============ 配置部分 ============
# 尝试从配置文件加载 API keys
try:
    from config_loader import load_config
    config = load_config()
    DASHSCOPE_KEY = config.get('DASHSCOPE_API_KEY') or os.getenv("DASHSCOPE_API_KEY") or "dummy-key"
    DEEPSEEK_KEY = config.get('DEEPSEEK_API_KEY') or os.getenv("DEEPSEEK_API_KEY") or "dummy-key"
    EMBED_MODEL = config.get('EMBED_MODEL', 'text-embedding-v4')
    LLM_MODEL = config.get('CHAT_MODEL', 'deepseek-chat')
except ImportError:
    # 如果没有配置文件，回退到环境变量或默认值
    DASHSCOPE_KEY = os.getenv("DASHSCOPE_API_KEY") or "dummy-key"
    DEEPSEEK_KEY = os.getenv("DEEPSEEK_API_KEY") or "dummy-key"
    EMBED_MODEL = "text-embedding-v4"
    LLM_MODEL = "deepseek-chat"

# 通义千问（DashScope）用于 Embedding；DeepSeek 用于 Chat
# 打包时使用 dummy key，运行时会从 config.json 重新加载
emb_client = OpenAI(api_key=DASHSCOPE_KEY, base_url="https://dashscope.aliyuncs.com/compatible-mode/v1")
chat_client = OpenAI(api_key=DEEPSEEK_KEY, base_url="https://api.deepseek.com/v1")
# 使用相对于脚本文件的路径，避免中文路径问题
SCRIPT_DIR = Path(__file__).parent.absolute()
TEST_DIR = SCRIPT_DIR / "相似文件测试/txt测试" # 测试文件夹路径
SIMILARITY_THRESHOLD = 0.6  # 相似度阈值，超过此值认为文件相似（降低阈值以便分析更多文件）
PDF_TEXT_MIN_LENGTH = 80    # 判断 PDF 是否为文本型的最小字符数
MAX_PDF_OCR_PAGES = None    # OCR 时最多处理的页数，None 表示处理所有页
IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".bmp", ".tif", ".tiff"}
IMAGE_ANALYSIS_CACHE: Dict[str, Dict[str, Any]] = {}
PHOTO_FEATURE_CACHE: Dict[str, np.ndarray] = {}
PHOTO_PHASH_CACHE: Dict[str, np.ndarray] = {}
PHOTO_HIST_CACHE: Dict[str, np.ndarray] = {}

# ============ 图片分类常量 ============
# 非照片类四分类（按删除可能性从高到低）
CATEGORY_CN = {
    "temporary": "临时类",
    "reference": "参考类",
    "saved": "收藏类",
    "memory": "记忆类"
}

# 删除可能性权重（0-1，越高越容易删除）
CATEGORY_DELETE_WEIGHT = {
    "temporary": 0.9,   # 90% 删除倾向
    "reference": 0.6,   # 60% 删除倾向
    "saved": 0.3,       # 30% 删除倾向
    "memory": 0.1       # 10% 删除倾向
}


def compute_photo_feature(path: str) -> np.ndarray:
    """将照片转换为低维特征向量，用于相似度比较"""
    if Image is None:
        return None
    if path in PHOTO_FEATURE_CACHE:
        return PHOTO_FEATURE_CACHE[path]
    try:
        with Image.open(path) as img:
            img = img.convert('RGB').resize((32, 32))
            arr = np.asarray(img, dtype=np.float32).flatten()
            norm = np.linalg.norm(arr)
            if norm > 0:
                arr = arr / norm
            PHOTO_FEATURE_CACHE[path] = arr
            return arr
    except Exception as e:
        print(f"⚠️ 生成照片特征失败：{path} -> {e}")
        return None


def photo_cosine_similarity(f1: np.ndarray, f2: np.ndarray) -> float:
    if f1 is None or f2 is None:
        return -1.0
    if f1.shape != f2.shape:
        return -1.0
    return float(np.dot(f1, f2) / (np.linalg.norm(f1) * np.linalg.norm(f2) + 1e-10))


def embedding_cosine_similarity(e1: Any, e2: Any) -> float:
    if e1 is None or e2 is None:
        return -1.0
    try:
        a = np.asarray(e1, dtype=float)
        b = np.asarray(e2, dtype=float)
    except Exception:
        return -1.0
    if a.shape != b.shape:
        return -1.0
    return float(np.dot(a, b) / (np.linalg.norm(a) * np.linalg.norm(b) + 1e-10))


def compute_phash_vector(path: str) -> np.ndarray:
    """简单 pHash，实现为 8x8 DCT"""
    if Image is None:
        return None
    if path in PHOTO_PHASH_CACHE:
        return PHOTO_PHASH_CACHE[path]
    try:
        with Image.open(path) as img:
            img = img.convert('L').resize((32, 32))
            pixels = np.asarray(img, dtype=np.float32)
            dct = np.fft.fft2(pixels)
            dct_low = np.abs(dct[:8, :8])
            med = np.median(dct_low)
            phash = (dct_low > med).astype(np.uint8).flatten()
            PHOTO_PHASH_CACHE[path] = phash
            return phash
    except Exception as e:
        print(f"⚠️ 生成 pHash 失败：{path} -> {e}")
        return None


def phash_similarity(a: np.ndarray, b: np.ndarray) -> float:
    if a is None or b is None or a.shape != b.shape:
        return 0.0
    same = np.sum(a == b)
    return same / len(a)


def compute_hsv_hist(path: str, bins: int = 16) -> np.ndarray:
    if Image is None:
        return None
    if path in PHOTO_HIST_CACHE:
        return PHOTO_HIST_CACHE[path]
    try:
        with Image.open(path) as img:
            hsv = img.convert('HSV')
            arr = np.asarray(hsv, dtype=np.float32)
            hist = []
            for channel in range(3):
                h, _ = np.histogram(arr[..., channel], bins=bins, range=(0, 255))
                hist.append(h.astype(np.float32))
            hist = np.concatenate(hist)
            norm = np.linalg.norm(hist)
            if norm > 0:
                hist = hist / norm
            PHOTO_HIST_CACHE[path] = hist
            return hist
    except Exception as e:
        print(f"⚠️ 生成 HSV 直方图失败：{path} -> {e}")
        return None


def histogram_similarity(a: np.ndarray, b: np.ndarray) -> float:
    if a is None or b is None or a.shape != b.shape:
        return 0.0
    return float(np.dot(a, b) / (np.linalg.norm(a) * np.linalg.norm(b) + 1e-10))

# Embedding 分块处理配置
EMBEDDING_MAX_TOKENS = 512       # 嵌入模型的最大 token 数（设置为512 tokens）
EMBEDDING_CHUNK_SIZE = 1024      # 每个分块的最大字符数（约1024字符对应512 tokens，中文1字符≈1-2 tokens）
EMBEDDING_CHUNK_OVERLAP = 100    # 分块重叠字符数，保持上下文连续性（减小重叠以适应更小的分块）

# OCR 配置
OCR_MODEL = "qwen3-vl-plus"      # 在线OCR模型
PDF_RENDER_SCALE = 300 / 72      # PDF 渲染分辨率（DPI），300（平衡，推荐）

# 初始化 ImageAnalyzer
image_analyzer = None
if ImageAnalyzer:
    try:
        # 复用已有的客户端
        # emb_client (DashScope) 既用于 embedding 也用于 VL (qwen-vl)
        # chat_client (DeepSeek) 用于逻辑分析
        image_analyzer = ImageAnalyzer(
            emb_client=emb_client,
            vl_client=emb_client,  # DashScope 支持 Qwen-VL
            llm_client=chat_client, # DeepSeek
            vl_model=OCR_MODEL,    # 复用配置的模型名
            debug=True             # 开启调试模式，显示 embedding 错误
        )
        print("✅ ImageAnalyzer 初始化成功")
    except Exception as e:
        print(f"⚠️ ImageAnalyzer 初始化失败: {e}")

# 文字层面相似度阈值（用于初步筛选，避免对完全不相似的文件进行embedding计算）
TEXT_SIMILARITY_THRESHOLD = 0.3  # 文字层面相似度阈值，低于此值直接跳过embedding比较


# ============ PDF / OCR 支持 ============
def extract_text_from_pdf(pdf_path: str) -> str:
    """优先使用 PyPDF2 提取文本型 PDF 内容（处理所有页面）"""
    if PdfReader is None:
        print("⚠️ 未安装 PyPDF2，无法直接解析 PDF。")
        return ""

    try:
        reader = PdfReader(pdf_path)
        total_pages = len(reader.pages)
        text_chunks = []
        
        print(f"ℹ️ 开始提取 PDF 文本（共 {total_pages} 页）：{Path(pdf_path).name}")
        
        # 显式遍历所有页面，确保处理所有页面
        for page_num in range(total_pages):
            try:
                page = reader.pages[page_num]
                page_text = page.extract_text() or ""
                if page_text:
                    text_chunks.append(page_text)
            except Exception as e:
                print(f"   ⚠️ 提取第 {page_num + 1} 页文本失败: {e}")
                continue
        
        extracted_pages = len(text_chunks)
        total_chars = sum(len(chunk) for chunk in text_chunks)
        result = "\n".join(text_chunks).strip()
        
        print(f"✅ PDF 文本提取完成：处理了 {extracted_pages}/{total_pages} 页，提取 {total_chars} 字符")
        
        return result
    except Exception as e:
        print(f"⚠️ 读取 PDF 文本失败 {pdf_path}: {e}")
        import traceback
        print(f"   详细错误: {traceback.format_exc()}")
        return ""


def pil_image_to_base64(pil_img) -> str:
    """
    将PIL Image转换为base64编码的图片URL格式
    
    Args:
        pil_img: PIL Image 对象
    
    Returns:
        base64编码的图片URL字符串
    """
    if pil_img is None:
        return ""
    
    try:
        # 将PIL Image转换为RGB模式（如果需要）
        if pil_img.mode != 'RGB':
            pil_img = pil_img.convert('RGB')
        
        # 将图片保存到BytesIO
        buffer = BytesIO()
        pil_img.save(buffer, format='JPEG', quality=95)
        buffer.seek(0)
        
        # 转换为base64
        image_bytes = buffer.read()
        base64_str = base64.b64encode(image_bytes).decode('utf-8')
        
        # 返回data URL格式
        return f"data:image/jpeg;base64,{base64_str}"
    except Exception as e:
        print(f"⚠️ 图片转换base64失败: {e}")
        return ""


def ocr_image_with_api(pil_img, prompt_text="请识别并提取图片中的所有文字内容，保持原有的格式和段落结构。") -> str:
    """
    使用在线OCR API（qwen3-vl-plus）识别图片中的文字
    
    Args:
        pil_img: PIL Image 对象
        prompt_text: OCR提示词，指导模型如何识别文字
    
    Returns:
        识别出的文本内容
    """
    if pil_img is None:
        print("⚠️ 图片对象为空，无法执行 OCR。")
        return ""
    
    try:
        # 将PIL Image转换为base64
        image_url = pil_image_to_base64(pil_img)
        if not image_url:
            print("⚠️ 图片转换失败，无法执行 OCR。")
            return ""
        
        # 调用在线OCR API
        completion = emb_client.chat.completions.create(
            model=OCR_MODEL,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": image_url
                            },
                        },
                        {"type": "text", "text": prompt_text},
                    ],
                },
            ],
            stream=False,
        )
        
        # 提取识别结果
        raw_text = completion.choices[0].message.content.strip()
        return raw_text
        
    except Exception as e:
        print(f"⚠️ 在线OCR识别失败: {e}")
        import traceback
        print(f"   详细错误: {traceback.format_exc()}")
        return ""


def extract_text_from_pdf_via_ocr(pdf_path: str) -> str:
    """使用 pypdfium2 将 PDF 渲染为图片后进行 OCR"""
    if pdfium is None:
        print("⚠️ 未安装 pypdfium2，无法对图片型 PDF 执行 OCR。")
        return ""

    try:
        pdf = pdfium.PdfDocument(pdf_path)
    except Exception as e:
        print(f"⚠️ 打开 PDF 失败 {pdf_path}: {e}")
        return ""

    text_parts = []
    try:
        total_pages = len(pdf)
        # 确定要处理的页数
        pages_to_process = total_pages if MAX_PDF_OCR_PAGES is None else min(total_pages, MAX_PDF_OCR_PAGES)
        
        # 输出OCR处理开始信息
        if MAX_PDF_OCR_PAGES is not None and total_pages > MAX_PDF_OCR_PAGES:
            print(f"ℹ️ 使用 {OCR_MODEL} 模型进行 OCR 处理（共 {total_pages} 页，将处理前 {MAX_PDF_OCR_PAGES} 页）：{Path(pdf_path).name}")
        else:
            print(f"ℹ️ 使用 {OCR_MODEL} 模型进行 OCR 处理（共 {total_pages} 页）：{Path(pdf_path).name}")
        
        for index in range(pages_to_process):
            try:
                # 尝试不同的页面访问方式
                try:
                    page = pdf[index]  # 新版本 API
                except (TypeError, AttributeError):
                    page = pdf.get_page(index)  # 旧版本 API
                
                # 尝试不同的渲染方法
                pil_image = None
                try:
                    # 方法1: render() 方法返回 PdfBitmap，需要转换为 PIL Image
                    # 使用更高的分辨率提高 OCR 质量
                    bitmap = page.render(scale=PDF_RENDER_SCALE)
                    if bitmap:
                        # 检查 bitmap 类型并转换
                        if hasattr(bitmap, 'to_pil'):
                            # pypdfium2 新版本：使用 to_pil() 方法
                            pil_image = bitmap.to_pil()
                        elif hasattr(bitmap, 'asarray'):
                            # 通过 numpy 数组转换
                            import numpy as np
                            array = bitmap.asarray()
                            if Image:
                                pil_image = Image.fromarray(array)
                        elif hasattr(bitmap, 'convert'):
                            # 如果已经是 PIL Image（某些版本）
                            pil_image = bitmap
                        else:
                            # 尝试其他可能的转换方法
                            try:
                                # 某些版本可能支持直接转换
                                if Image:
                                    pil_image = Image.fromarray(np.array(bitmap))
                            except:
                                pass
                except Exception as render_error:
                    try:
                        # 方法2: render_topil() 方法（某些版本直接返回 PIL Image）
                        pil_image = page.render_topil(scale=PDF_RENDER_SCALE)
                    except AttributeError:
                        try:
                            # 方法3: render_to() 方法
                            from io import BytesIO
                            buffer = BytesIO()
                            page.render_to(buffer, scale=PDF_RENDER_SCALE)
                            buffer.seek(0)
                            if Image:
                                pil_image = Image.open(buffer)
                        except Exception as e:
                            print(f"   ⚠️ 渲染方法失败: {render_error}, 备选方法也失败: {e}")
                            pass
                
                if pil_image:
                    # 使用在线OCR API识别页面文字
                    page_text = ocr_image_with_api(
                        pil_image, 
                        prompt_text="请识别并提取PDF页面中的所有文字内容，保持原有的格式、段落结构和排版。如果是表格，请以Markdown表格格式输出。"
                    )
                    
                    if page_text:
                        text_parts.append(f"【第 {index + 1} 页】\n{page_text}")
                else:
                    print(f"   ⚠️ 第 {index + 1} 页渲染失败，无法获取图片对象")
                
                # 尝试关闭页面
                try:
                    page.close()
                except:
                    pass
                    
            except Exception as e:
                print(f"⚠️ 处理第 {index + 1}/{pages_to_process} 页时出错: {e}")
                continue
            
            # 显示进度
            if (index + 1) % 10 == 0 or (index + 1) == pages_to_process:
                print(f"   📄 已处理 {index + 1}/{pages_to_process} 页...")
        
        # 计算提取的字符数
        total_chars = sum(len(t) for t in text_parts)
        extracted_pages = len(text_parts)
        
        print(f"✅ 使用 {OCR_MODEL} 模型 OCR 处理完成：处理了 {extracted_pages}/{pages_to_process} 页，提取 {total_chars} 字符")
        return "\n\n".join(text_parts).strip()
    finally:
        try:
            pdf.close()
        except:
            pass


def extract_text_from_image(image_path: str) -> str:
    """对单张图片执行 OCR（使用在线OCR API）"""
    if Image is None:
        print("⚠️ 未安装 Pillow，无法读取图片。")
        return ""

    try:
        print(f"ℹ️ 使用 {OCR_MODEL} 模型进行 OCR 处理：{Path(image_path).name}")
        with Image.open(image_path) as pil_img:
            ocr_text = ocr_image_with_api(
                pil_img,
                prompt_text="请识别并提取图片中的所有文字内容，保持原有的格式和段落结构。如果是表格，请以Markdown表格格式输出。"
            )
            if ocr_text:
                print(f"✅ 使用 {OCR_MODEL} 模型 OCR 处理完成：提取 {len(ocr_text)} 字符")
            return ocr_text
    except Exception as e:
        print(f"⚠️ 打开图片失败 {image_path}: {e}")
        return ""

# ============ 工具函数部分 ============
def split_text_into_chunks(text: str, chunk_size: int = None, overlap: int = None) -> List[str]:
    """
    将长文本分割成多个块，用于分批处理 embedding
    
    Args:
        text: 要分割的文本
        chunk_size: 每个分块的最大字符数，None 时使用默认配置
        overlap: 分块重叠字符数，None 时使用默认配置
    
    Returns:
        文本块列表
    """
    if not text:
        return []
    
    chunk_size = chunk_size or EMBEDDING_CHUNK_SIZE
    overlap = overlap or EMBEDDING_CHUNK_OVERLAP
    
    # 如果文本长度小于分块大小，直接返回
    if len(text) <= chunk_size:
        return [text]
    
    chunks = []
    start = 0
    
    while start < len(text):
        # 计算当前块的结束位置
        end = start + chunk_size
        
        # 如果还没到文本末尾，尝试在合适的位置断开（优先在换行符或句号处）
        if end < len(text):
            # 向前查找换行符或句号
            break_pos = end
            for i in range(end, max(start + chunk_size - 500, start), -1):
                if text[i] in ['\n', '。', '.', '！', '!', '？', '?']:
                    break_pos = i + 1
                    break
            
            # 如果没找到合适的断点，使用原位置
            if break_pos == end:
                # 尝试在空格处断开
                for i in range(end, max(start + chunk_size - 200, start), -1):
                    if text[i] in [' ', '\t']:
                        break_pos = i + 1
                        break
            
            end = break_pos
        
        # 提取当前块
        chunk = text[start:end]
        chunks.append(chunk)
        
        # 计算下一个块的起始位置（考虑重叠）
        start = end - overlap
        if start >= len(text):
            break
    
    return chunks


def get_embedding(text: str, use_chunking: bool = True) -> List[float]:
    """
    调用通义千问文本嵌入 API，支持长文本分块处理
    
    Args:
        text: 要嵌入的文本
        use_chunking: 是否对长文本进行分块处理
    
    Returns:
        嵌入向量（单个向量或平均向量）
    """
    if not text:
        return None
    
    try:
        # 如果文本较短或不需要分块，直接处理
        if not use_chunking or len(text) <= EMBEDDING_CHUNK_SIZE:
            response = emb_client.embeddings.create(model=EMBED_MODEL, input=text)
            return response.data[0].embedding
        
        # 长文本需要分块处理
        chunks = split_text_into_chunks(text)
        
        if len(chunks) == 1:
            # 只有一个块，直接处理
            response = emb_client.embeddings.create(model=EMBED_MODEL, input=chunks[0])
            return response.data[0].embedding
        
        # 多个块，分别获取 embedding 然后加权平均
        total_chars_in_chunks = sum(len(chunk) for chunk in chunks)
        print(f"   📦 文本较长（{len(text)} 字符），分 {len(chunks)} 块处理（总字符数验证: {total_chars_in_chunks} 字符）...")
        
        # 验证分块完整性
        if abs(total_chars_in_chunks - len(text)) > len(text) * 0.05:  # 允许5%的误差（由于重叠）
            print(f"   ⚠️ 警告：分块字符总数 ({total_chars_in_chunks}) 与原文 ({len(text)}) 差异较大，可能存在文本丢失")
        
        embeddings = []
        chunk_weights = []  # 存储每个块的权重（按字符数）
        
        for i, chunk in enumerate(chunks, 1):
            try:
                response = emb_client.embeddings.create(model=EMBED_MODEL, input=chunk)
                embedding = response.data[0].embedding
                embeddings.append(embedding)
                chunk_weights.append(len(chunk))  # 使用字符数作为权重
                
                # 显示进度
                if len(chunks) > 5 and i % 5 == 0:
                    print(f"     已处理 {i}/{len(chunks)} 块（当前块 {len(chunk)} 字符）...")
                    
            except Exception as e:
                print(f"   ⚠️ 第 {i} 块 embedding 失败: {e}，该块字符数: {len(chunk)}")
                continue
        
        if not embeddings:
            print(f"   ❌ 所有分块的 embedding 都失败了")
            return None
        
        # 验证处理的块数
        if len(embeddings) < len(chunks):
            print(f"   ⚠️ 警告：只成功处理了 {len(embeddings)}/{len(chunks)} 个分块，部分文本可能未参与embedding计算")
        
        # 计算加权平均 embedding（按块长度加权，确保长块有更大权重）
        if len(embeddings) == 1:
            return embeddings[0]
        
        # 转换为 numpy 数组进行加权平均
        embeddings_array = np.array(embeddings)
        weights_array = np.array(chunk_weights)
        # 归一化权重
        weights_array = weights_array / weights_array.sum()
        # 加权平均
        avg_embedding = np.average(embeddings_array, axis=0, weights=weights_array).tolist()
        
        processed_chars = sum(chunk_weights)
        print(f"   ✅ 成功合并 {len(embeddings)} 个分块的 embedding（加权平均，处理了 {processed_chars}/{len(text)} 字符）")
        return avg_embedding
        
    except Exception as e:
        err_msg = str(e)
        print(f"❌ 获取嵌入失败（{EMBED_MODEL}）: {err_msg}")
        return None


def text_similarity(text1: str, text2: str) -> float:
    """
    计算两个文本在文字层面的相似度（用于初步筛选）
    
    使用多种方法组合：
    1. Jaccard相似度（基于字符集合）
    2. 字符重叠比例
    3. 长度相似度
    
    Args:
        text1: 第一个文本
        text2: 第二个文本
    
    Returns:
        相似度分数（0-1之间）
    """
    if not text1 or not text2:
        return 0.0
    
    # 方法1: Jaccard相似度（基于字符集合）
    set1 = set(text1)
    set2 = set(text2)
    if len(set1) == 0 and len(set2) == 0:
        jaccard = 1.0
    elif len(set1) == 0 or len(set2) == 0:
        jaccard = 0.0
    else:
        intersection = len(set1 & set2)
        union = len(set1 | set2)
        jaccard = intersection / union if union > 0 else 0.0
    
    # 方法2: 字符重叠比例（考虑字符频率）
    # 计算公共字符的总数（考虑重复）
    common_chars = 0
    text1_chars = {}
    text2_chars = {}
    
    for char in text1:
        text1_chars[char] = text1_chars.get(char, 0) + 1
    for char in text2:
        text2_chars[char] = text2_chars.get(char, 0) + 1
    
    for char in set(text1) & set(text2):
        common_chars += min(text1_chars[char], text2_chars[char])
    
    total_chars = len(text1) + len(text2)
    overlap_ratio = (2 * common_chars) / total_chars if total_chars > 0 else 0.0
    
    # 方法3: 长度相似度
    len1, len2 = len(text1), len(text2)
    if len1 == 0 and len2 == 0:
        len_sim = 1.0
    elif len1 == 0 or len2 == 0:
        len_sim = 0.0
    else:
        len_sim = min(len1, len2) / max(len1, len2)
    
    # 综合相似度（加权平均）
    # Jaccard权重0.4，重叠比例权重0.4，长度相似度权重0.2
    similarity = 0.4 * jaccard + 0.4 * overlap_ratio + 0.2 * len_sim
    
    return similarity


def cosine_similarity(a: List[float], b: List[float]) -> float:
    """计算余弦相似度"""
    if a is None or b is None:
        return -1.0
    a, b = np.array(a), np.array(b)
    if a.shape != b.shape:
        return -1.0
    return np.dot(a, b) / (np.linalg.norm(a) * np.linalg.norm(b) + 1e-10)


def read_file_content(file_path: str) -> str:
    """
    读取文件内容，支持 txt、md、docx、pdf、图片等文件。
    对于不支持的文件类型，将返回空字符串并打印提示。
    """
    ext = Path(file_path).suffix.lower()
    try:
        if ext == ".txt" or ext == ".md":
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        elif ext == ".docx":
            doc = Document(file_path)
            return "\n".join([p.text for p in doc.paragraphs])
        elif ext == ".pdf":
            # 先尝试直接提取文本
            text_content = extract_text_from_pdf(file_path)
            if len(text_content) >= PDF_TEXT_MIN_LENGTH:
                print(f"✅ PDF 文本提取成功（{len(text_content)} 字符）：{Path(file_path).name}")
                return text_content

            # 如果文本提取失败或内容太少，尝试 OCR
            print(f"ℹ️ PDF 文本提取内容较少（{len(text_content)} 字符），尝试 OCR：{Path(file_path).name}")
            ocr_content = extract_text_from_pdf_via_ocr(file_path)
            if ocr_content:
                print(f"✅ PDF OCR 成功（{len(ocr_content)} 字符）：{Path(file_path).name}")
                return ocr_content

            # 如果 OCR 也失败，返回原始提取结果（如果有）
            if text_content:
                print(f"⚠️ PDF OCR 失败，返回原始提取结果（{len(text_content)} 字符）：{Path(file_path).name}")
                return text_content
            else:
                print(f"❌ 无法从 PDF 提取文本（文本提取和 OCR 均失败）：{Path(file_path).name}")
                return ""
        elif ext in IMAGE_EXTENSIONS:
            if image_analyzer is None:
                print(f"⚠️ 未初始化 ImageAnalyzer，跳过图片视觉理解：{file_path}")
                return ""
            # 对于图片，返回更详细的描述用于 embedding
            analysis = get_or_run_image_analysis(file_path)
            if analysis:
                # 构建更丰富的文本描述
                parts = []
                # 1. 短描述（最重要）
                if analysis.get('short_description'):
                    parts.append(analysis['short_description'])
                # 2. 场景描述
                vl = analysis.get('vl', {})
                if vl.get('scene'):
                    parts.append(vl['scene'])
                # 3. 人脸信息
                if vl.get('faces') and vl['faces'] > 0:
                    parts.append(f"{vl['faces']} people in image")
                    if vl.get('eyes_open') is not None:
                        parts.append("eyes open" if vl['eyes_open'] else "eyes closed")
                # 4. 图片类型（photo 很重要）
                if vl.get('type'):
                    parts.append(f"type: {vl['type']}")
                
                return " | ".join(parts) if parts else Path(file_path).stem
            print(f"⚠️ 图片视觉理解失败：{file_path}")
            return ""
        else:
            print(f"⚠️ 不支持的文件类型：{ext}")
            return ""
    except Exception as e:
        print(f"⚠️ 读取文件失败 {file_path}: {e}")
        return ""


def scan_files(directory) -> List[Dict]:
    """扫描目录下的所有文件，返回文件信息列表"""
    file_infos = []
    directory_path = Path(directory) if isinstance(directory, (str, Path)) else directory
    
    if not directory_path.exists():
        print(f"❌ 目录不存在: {directory}")
        return file_infos
    
    # 排序文件列表以确保跨平台一致性
    for file_path in sorted(directory_path.iterdir(), key=lambda p: p.name.lower()):
        if file_path.is_file():
            file_info = {
                "path": str(file_path),
                "name": file_path.name,
                "content": read_file_content(str(file_path))
            }
            
            # 对于图片，直接附加 embedding（ImageAnalyzer 已经生成了）
            ext = file_path.suffix.lower()
            if ext in IMAGE_EXTENSIONS and image_analyzer is not None:
                analysis = get_or_run_image_analysis(str(file_path))
                if analysis and analysis.get('embedding'):
                    file_info['embedding'] = analysis['embedding']
            
            file_infos.append(file_info)
    
    return file_infos


def find_similar_file_pairs(file_infos: List[Dict], similarity_threshold: float = 0.7) -> List[Tuple[Dict, Dict, float]]:
    """
    找到相似的文件对
    
    流程：
    1. 先进行文字层面的相似度比较（快速筛选）
    2. 对文字层面相似的文件进行embedding计算
    3. 使用embedding进行语义相似度比较
    """
    # 过滤掉没有内容的文件
    valid_files = [f for f in file_infos if f.get("content")]
    
    if len(valid_files) < 2:
        print("❌ 有效文件数量不足，无法进行相似度比较")
        return []
    
    print(f"\n📝 第一步：文字层面相似度比较（筛选候选文件对）...")
    text_similar_pairs = []
    text_similarities = []
    
    # 先进行文字层面的相似度比较
    for file1, file2 in tqdm(combinations(valid_files, 2), desc="文字层面比较", total=len(valid_files)*(len(valid_files)-1)//2):
        text_sim = text_similarity(file1["content"], file2["content"])
        text_similarities.append((file1["name"], file2["name"], text_sim))
        
        # 只保留文字层面相似度超过阈值的文件对
        if text_sim >= TEXT_SIMILARITY_THRESHOLD:
            text_similar_pairs.append((file1, file2, text_sim))
    
    print(f"   ✅ 文字层面比较完成，找到 {len(text_similar_pairs)} 对候选文件（文字相似度 >= {TEXT_SIMILARITY_THRESHOLD}）")
    
    if len(text_similar_pairs) == 0:
        print(f"\n✅ 未找到文字层面相似的文件对（文字相似度阈值: {TEXT_SIMILARITY_THRESHOLD}）")
        return []
    
    # 收集需要计算embedding的文件（去重）
    unique_files = {}
    for file1, file2, _ in text_similar_pairs:
        unique_files[id(file1)] = file1
        unique_files[id(file2)] = file2
    
    # 为候选文件生成嵌入向量
    print(f"\n🔍 第二步：为 {len(unique_files)} 个候选文件生成嵌入向量...")
    
    for file_id, file_info in tqdm(unique_files.items(), desc="生成嵌入向量"):
        if file_info.get("embedding") is None:
            file_info["embedding"] = get_embedding(file_info["content"])
    
    # 过滤掉embedding失败的文件对
    valid_pairs = []
    for file1, file2, text_sim in text_similar_pairs:
        if file1.get("embedding") is not None and file2.get("embedding") is not None:
            valid_pairs.append((file1, file2, text_sim))
    
    if len(valid_pairs) == 0:
        print("❌ 所有候选文件对的embedding生成都失败了")
        return []
    
    print(f"\n🧮 第三步：使用embedding进行语义相似度比较...")
    similar_pairs = []
    all_similarities = []
    
    # 使用embedding计算语义相似度
    for file1, file2, text_sim in tqdm(valid_pairs, desc="语义相似度比较"):
        semantic_sim = cosine_similarity(file1["embedding"], file2["embedding"])
        all_similarities.append((file1["name"], file2["name"], text_sim, semantic_sim))
        
        if semantic_sim >= similarity_threshold:
            similar_pairs.append((file1, file2, semantic_sim))
    
    # 输出所有文件对的相似度（文字层面 + 语义层面）
    print(f"\n📊 所有文件对的相似度（文字层面 | 语义层面）：")
    for name1, name2, text_sim, sem_sim in sorted(all_similarities, key=lambda x: x[3], reverse=True):
        print(f"   {name1} <-> {name2}: 文字={text_sim:.3f} | 语义={sem_sim:.3f}")
    
    # 按语义相似度降序排序
    similar_pairs.sort(key=lambda x: x[2], reverse=True)
    
    return similar_pairs


def analyze_and_report_images(file_infos: List[Dict], output_dir: Path):
    """
    对扫描到的图片进行深度分析（BestShot、内容理解、删除/保留建议）
    并将结果输出到控制台和日志文件
    """
    if image_analyzer is None:
        return

    images = [f for f in file_infos if Path(f['path']).suffix.lower() in IMAGE_EXTENSIONS]
    if not images:
        return

    print(f"\n🖼️ 正在对 {len(images)} 张图片进行深度分析与去重建议...")
    
    analysis_results = []
    
    # 准备详细日志文件
    log_file = output_dir / "image_analysis_details.txt"
    
    # 收集分析结果后再统一输出，便于执行相似照片去重
    for img_info in tqdm(images, desc="图片分析"):
        path = img_info['path']
        name = img_info['name']
        
        try:
            # 基础分析（VL模型）
            res = get_or_run_image_analysis(path)
            if not res:
                raise RuntimeError("无法获取图片分析结果")
            
            # 判断：真实照片 or 非照片类
            is_real_photo = not res.get('likely_screenshot', False)
            
            if is_real_photo:
                # 真实照片 → 提取特征用于 BestShot 去重
                res['photo_feature'] = compute_photo_feature(path)
            else:
                # 非照片类 → LLM 四分类
                llm_result = image_analyzer.screenshot_classify_with_llm(res)
                res['category'] = llm_result.get('category', 'saved')
                res['app_name'] = llm_result.get('app_name', '')
                
                category_cn = CATEGORY_CN.get(res['category'], '未知')
                suggestion_text = llm_result.get('suggestion', 'keep')
                print(f"      🏷️ 分类: {category_cn} | 建议: {suggestion_text}")
                
                # LLM 建议删除 → 标记
                if llm_result.get('suggestion') == 'delete':
                    res.setdefault('suggestion', {})
                    res['suggestion']['delete'] = True
                    res['suggestion']['reason'] = f'LLM建议删除（{category_cn}）'
            
            # 同步 embedding 到 img_info（用于相似度比较）
            if res.get('embedding'):
                img_info['embedding'] = res['embedding']
            
            analysis_results.append(res)

        except Exception as e:
            print(f"❌ 分析图片 {name} 失败: {e}")
    
    # 分类：真实照片 vs 非照片类
    real_photos = [r for r in analysis_results if not r.get('likely_screenshot')]
    non_photos = [r for r in analysis_results if r.get('likely_screenshot')]
    
    # 统计非照片类的分布
    category_count = {'temporary': 0, 'reference': 0, 'saved': 0, 'memory': 0}
    for r in non_photos:
        cat = r.get('category', 'saved')
        category_count[cat] = category_count.get(cat, 0) + 1
    
    print(f"\n📊 图片分类统计：")
    print(f"   📸 真实照片: {len(real_photos)} 张（BestShot 去重）")
    print(f"\n   非照片类：")
    print(f"      🗑️ 临时类: {category_count['temporary']} 张（删除可能性最高）")
    print(f"      📋 参考类: {category_count['reference']} 张（删除可能性中高）")
    print(f"      💾 收藏类: {category_count['saved']} 张（删除可能性中低）")
    print(f"      ❓ 记忆类: {category_count['memory']} 张（扫描件等）")
    
    # 去重
    mark_similar_photos(real_photos)
    mark_similar_screenshots(non_photos)
    
    print(f"\n{'='*60}")
    print(f"{'图片文件':<20} | {'类型':<12} | {'标签/描述':<18} | {'处理建议'}")
    print(f"{'-'*60}")
    for res in analysis_results:
        name = Path(res['path']).name
        # 显示类型
        if not res.get('likely_screenshot'):
            img_type = "📸 照片"
        else:
            img_type = CATEGORY_CN.get(res.get('category', 'saved'), '收藏类')
        
        # 描述信息
        desc = res.get('short_description', '') or res.get('vl', {}).get('scene', '')
        if desc and len(desc) > 18:
            desc = desc[:15] + "..."
        
        suggestion = format_image_suggestion(res)
        print(f"{name:<20} | {img_type:<12} | {desc:<18} | {suggestion}")
    print(f"{'='*60}")
    print_screenshot_tree(analysis_results)
    
    # --- 保存详细报告 ---
    try:
        with open(log_file, 'w', encoding='utf-8') as f:
            f.write("图片深度分析与去重详细报告\n")
            f.write(f"生成时间: {datetime.now()}\n")
            f.write("="*60 + "\n\n")
            
            for res in analysis_results:
                name = Path(res['path']).name
                f.write(f"文件名: {name}\n")
                f.write(f"路径: {res['path']}\n")
                
                # 分类信息
                if not res.get('likely_screenshot'):
                    # 真实照片
                    f.write(f"类型: 真实照片（BestShot 去重）\n")
                    score = res.get('bestshot_score', 0)
                    f.write(f"质量分数: {score:.2f}\n")
                else:
                    # 非照片类
                    category = res.get('category', 'saved')
                    category_cn = CATEGORY_CN.get(category, '收藏类')
                    f.write(f"类型: {category_cn}\n")
                    if res.get('app_name'):
                        f.write(f"APP: {res['app_name']}\n")
                    delete_weight = CATEGORY_DELETE_WEIGHT.get(category, 0.3)
                    f.write(f"删除倾向: {delete_weight*100:.0f}%\n")
                
                # 视觉描述
                desc = res.get('short_description') or res.get('vl', {}).get('scene', '')
                if desc:
                    f.write(f"描述: {desc}\n")
                
                # Delete Suggestion
                sug = res.get('suggestion', {})
                if sug.get('delete'):
                    f.write(f"⚠️ 删除建议: 建议删除。原因: {sug.get('reason')}\n")
                
                f.write("-" * 40 + "\n")
        
        print(f"✅ 详细分析报告已保存至: {log_file}")
            
    except Exception as e:
        print(f"⚠️ 保存详细报告失败: {e}")


def get_or_run_image_analysis(path: str) -> Dict[str, Any]:
    """获取缓存的图片分析结果，否则调用 ImageAnalyzer"""
    if image_analyzer is None:
        return {}
    if path in IMAGE_ANALYSIS_CACHE:
        return IMAGE_ANALYSIS_CACHE[path]
    try:
        res = image_analyzer.analyze_image_file(path)
        IMAGE_ANALYSIS_CACHE[path] = res
        return res
    except Exception as e:
        print(f"⚠️ 调用 ImageAnalyzer 失败：{path} -> {e}")
        return {}


def image_analysis_to_text(res: Dict[str, Any]) -> str:
    """将图片分析结果压缩为文本摘要（用于相似度比较）"""
    parts = []
    
    # 描述信息
    if res.get('short_description'):
        parts.append(res['short_description'])
    
    # 类型信息
    if not res.get('likely_screenshot'):
        parts.append("type:photo")
    else:
        category = res.get('category', 'saved')
        parts.append(f"type:{category}")
        if res.get('app_name'):
            parts.append(f"app:{res['app_name']}")
    
    return " | ".join(parts) if parts else Path(res.get('path', '')).stem


def mark_similar_photos(photo_items: List[Dict[str, Any]],
                        feature_high: float = 0.92,          #照片特征相似度阈值(高) - 降低到0.92以识别更多相似图片
                        feature_low: float = 0.88,           #照片特征相似度阈值(低) - 降低到0.88
                        embed_threshold: float = 0.85,       #embedding相似度阈值 - 降低到0.85
                        desc_threshold: float = 0.83,        #照片描述相似度阈值 - 降低到0.83
                        phash_high: float = 0.90,            #pHash高阈值 - 降低到0.90
                        phash_low: float = 0.85,             #pHash低阈值 - 降低到0.85
                        hist_threshold: float = 0.80,        #HSV直方图阈值 - 降低到0.80
                        debug: bool = True):                 #调试模式
    """对照片进行相似度聚类，仅保留每个簇中的最佳照片"""
    if len(photo_items) <= 1:
        return
    n = len(photo_items)
    features = []
    embeddings = []
    descriptions = []
    person_ids = []
    phashes = []
    hists = []
    
    if debug:
        print(f"\n🔍 照片相似度分析（共 {n} 张照片）")
        print(f"   阈值设置：")
        print(f"      - 照片特征: high={feature_high}, low={feature_low}")
        print(f"      - pHash: high={phash_high}, low={phash_low}")
        print(f"      - HSV直方图: {hist_threshold}")
        print(f"      - Embedding: {embed_threshold}")
        print(f"      - 描述相似度: {desc_threshold}")
    
    # 先收集所有数据
    for idx, res in enumerate(photo_items):
        feat = res.get('photo_feature')
        if feat is None:
            feat = compute_photo_feature(res['path'])
            res['photo_feature'] = feat
        features.append(feat)
        
        emb = res.get('embedding')
        embeddings.append(emb)
        
        # 调试：检查每张照片的 embedding
        if debug and idx < 3:  # 打印前3张的详细信息
            name = Path(res['path']).name
            print(f"\n      📋 照片 {idx+1}: {name}")
            print(f"         有 embedding: {emb is not None}")
            if emb is not None:
                print(f"         embedding 类型: {type(emb)}, 长度: {len(emb) if hasattr(emb, '__len__') else 'N/A'}")
            else:
                print(f"         res 中的所有键: {list(res.keys())}")
                if 'embedding' in res:
                    print(f"         embedding 值为: {res['embedding']}")
        
        descriptions.append(res.get('short_description') or res.get('vl', {}).get('scene', ''))
        person_ids.append(extract_person_identifier(res['path']))
        phashes.append(compute_phash_vector(res['path']))
        hists.append(compute_hsv_hist(res['path']))
    
    # 统计数据状态（在数据收集完成后）
    if debug:
        emb_count = sum(1 for e in embeddings if e is not None)
        feat_count = sum(1 for f in features if f is not None)
        print(f"\n   数据状态：")
        print(f"      - Embedding 可用: {emb_count}/{n} 张")
        print(f"      - 照片特征可用: {feat_count}/{n} 张")
        if emb_count < n:
            print(f"      ⚠️ 警告：有 {n - emb_count} 张照片缺少 embedding，可能影响相似度判断")
    
    parent = list(range(n))

    def find(x):
        while parent[x] != x:
            parent[x] = parent[parent[x]]
            x = parent[x]
        return x

    def union(a, b):
        ra, rb = find(a), find(b)
        if ra != rb:
            parent[rb] = ra

    similarity_details = []  # 存储相似度详情用于调试
    
    for i in range(n):
        fi = features[i]
        if fi is None:
            continue
        for j in range(i + 1, n):
            fj = features[j]
            if fj is None:
                continue
            if should_force_keep(person_ids[i], person_ids[j]):
                continue
            
            # 计算所有维度的相似度
            photo_sim = photo_cosine_similarity(fi, fj)
            ph_sim = phash_similarity(phashes[i], phashes[j])
            hist_sim = histogram_similarity(hists[i], hists[j])
            embed_sim = embedding_cosine_similarity(embeddings[i], embeddings[j])
            desc_sim = short_desc_similarity(descriptions[i], descriptions[j])
            
            # 存储详情
            if debug:
                name_i = Path(photo_items[i]['path']).name
                name_j = Path(photo_items[j]['path']).name
                similarity_details.append({
                    'pair': (name_i, name_j),
                    'photo': photo_sim,
                    'phash': ph_sim,
                    'hist': hist_sim,
                    'embed': embed_sim,
                    'desc': desc_sim
                })
            
            # 快速过滤：如果 pHash 和照片特征都太低，跳过
            if ph_sim < phash_low and photo_sim < feature_low:
                continue
            
            # 多维度判断相似性
            similar = False
            match_reason = ""
            
            if ph_sim >= phash_high and hist_sim >= hist_threshold:
                similar = True
                match_reason = f"pHash高({ph_sim:.3f}) + 直方图({hist_sim:.3f})"
            elif photo_sim >= feature_high and hist_sim >= hist_threshold:
                similar = True
                match_reason = f"特征高({photo_sim:.3f}) + 直方图({hist_sim:.3f})"
            elif ph_sim >= phash_low and hist_sim >= hist_threshold and (photo_sim >= feature_low or embed_sim >= embed_threshold):
                similar = True
                match_reason = f"pHash中({ph_sim:.3f}) + 直方图({hist_sim:.3f}) + 特征/嵌入"
            elif embed_sim >= embed_threshold:
                similar = True
                match_reason = f"嵌入向量({embed_sim:.3f})"
            elif desc_sim >= desc_threshold:
                similar = True
                match_reason = f"描述相似({desc_sim:.3f})"
            
            if similar:
                union(i, j)
                if debug:
                    similarity_details[-1]['matched'] = True
                    similarity_details[-1]['reason'] = match_reason
            elif debug:
                similarity_details[-1]['matched'] = False
                similarity_details[-1]['reason'] = "未达阈值"

    groups: Dict[int, List[Dict[str, Any]]] = {}
    for idx in range(n):
        root = find(idx)
        groups.setdefault(root, []).append(photo_items[idx])
    
    # 调试输出：显示相似度详情
    if debug and similarity_details:
        print(f"\n📊 照片对相似度详情（共 {len(similarity_details)} 对）：")
        # 按匹配状态分组显示
        matched = [d for d in similarity_details if d.get('matched')]
        unmatched = [d for d in similarity_details if not d.get('matched')]
        
        if matched:
            print(f"\n   ✅ 匹配的照片对 ({len(matched)} 对)：")
            for detail in matched[:10]:  # 只显示前10对
                name_i, name_j = detail['pair']
                print(f"      {name_i} <-> {name_j}")
                print(f"         特征={detail['photo']:.3f} | pHash={detail['phash']:.3f} | "
                      f"直方图={detail['hist']:.3f} | 嵌入={detail['embed']:.3f} | 描述={detail['desc']:.3f}")
                print(f"         匹配原因: {detail['reason']}")
            if len(matched) > 10:
                print(f"      ... 还有 {len(matched) - 10} 对匹配的照片")
        
        if unmatched:
            print(f"\n   ❌ 未匹配的照片对 ({len(unmatched)} 对) - 显示前10对：")
            for detail in unmatched[:10]:
                name_i, name_j = detail['pair']
                print(f"      {name_i} <-> {name_j}")
                print(f"         特征={detail['photo']:.3f} | pHash={detail['phash']:.3f} | "
                      f"直方图={detail['hist']:.3f} | 嵌入={detail['embed']:.3f} | 描述={detail['desc']:.3f}")
                # 找出哪些维度接近阈值
                close_to = []
                if detail['photo'] >= feature_low * 0.8:
                    close_to.append(f"特征({detail['photo']:.3f}/{feature_low})")
                if detail['phash'] >= phash_low * 0.8:
                    close_to.append(f"pHash({detail['phash']:.3f}/{phash_low})")
                if detail['hist'] >= hist_threshold * 0.8:
                    close_to.append(f"直方图({detail['hist']:.3f}/{hist_threshold})")
                if detail['embed'] >= embed_threshold * 0.8:
                    close_to.append(f"嵌入({detail['embed']:.3f}/{embed_threshold})")
                if detail['desc'] >= desc_threshold * 0.8:
                    close_to.append(f"描述({detail['desc']:.3f}/{desc_threshold})")
                if close_to:
                    print(f"         接近阈值: {', '.join(close_to)}")

    for group in groups.values():
        if len(group) < 2:
            continue
        best = max(group, key=lambda r: r.get('bestshot_score', 0))
        for item in group:
            if item is best:
                continue
            item.setdefault('suggestion', {})
            # 只在尚未被其他规则标记删除时设置为删除
            if not item['suggestion'].get('delete'):
                item['suggestion']['delete'] = True
                item['suggestion']['reason'] = "相似照片，仅保留最佳镜头"
            item['duplicate_of'] = best['path']
    
    if debug:
        num_groups = len([g for g in groups.values() if len(g) >= 2])
        print(f"\n✅ 照片聚类完成：找到 {num_groups} 个相似组")


def format_image_suggestion(res: Dict[str, Any]) -> str:
    """格式化图片处理建议"""
    suggestion = res.get('suggestion', {})
    
    # 如果标记为删除
    if suggestion.get('delete'):
        if res.get('duplicate_of'):
            return f"🗑 删除（相似，保留 {Path(res['duplicate_of']).name}）"
        return f"🗑 删除（{suggestion.get('reason', '建议删除')}）"
    
    # 保留：显示质量或分类信息
    if not res.get('likely_screenshot'):
        # 真实照片
        score = res.get('bestshot_score')
        return f"✅ 保留（质量 {score:.2f}）" if score else "✅ 保留"
    else:
        # 非照片类
        category = res.get('category', 'saved')
        category_cn = CATEGORY_CN.get(category, '收藏类')
        app_name = res.get('app_name', '')
        detail = f"{category_cn} - {app_name}" if app_name else category_cn
        return f"✅ 保留（{detail}）"


def short_desc_similarity(a: str, b: str) -> float:
    if not a or not b:
        return 0.0
    ratio = SequenceMatcher(None, a, b).ratio()
    kw_sim = keyword_similarity(a, b)
    return 0.6 * ratio + 0.4 * kw_sim


def keyword_similarity(a: str, b: str) -> float:
    ka = extract_keywords(a)
    kb = extract_keywords(b)
    if not ka or not kb:
        return 0.0
    inter = ka & kb
    union = ka | kb
    return len(inter) / len(union)


def extract_keywords(text: str) -> set:
    tokens = re.findall(r"[A-Za-z]{3,}", text.lower())
    return set(tokens)


def is_cjk(char: str) -> bool:
    return '\u4e00' <= char <= '\u9fff'


def extract_person_identifier(path: str) -> str:
    stem = Path(path).stem
    tokens = re.split(r'[_\-\s]+', stem)
    for token in tokens:
        token = token.strip()
        if not token:
            continue
        cleaned = ''.join(ch for ch in token if ch.isalnum() or is_cjk(ch))
        if not cleaned:
            continue
        upper = cleaned.upper()
        if upper.startswith(('IMG', 'PXL', 'DSC', 'VID', 'PHOTO', 'SCREEN')):
            continue
        if any(is_cjk(ch) for ch in cleaned):
            name = ''.join(ch for ch in cleaned if is_cjk(ch))
            if name:
                return name
        if cleaned.isalpha() and len(cleaned) >= 3:
            return cleaned.lower()
    return ''


def should_force_keep(pid_a: str, pid_b: str) -> bool:
    return bool(pid_a and pid_b and pid_a != pid_b)


def compute_screenshot_clutter(res: Dict[str, Any]) -> float:
    text = (res.get('short_description') or '') + ' ' + (res.get('vl', {}).get('scene') or '')
    text = text.lower()
    keywords = [
        'video call', 'chat', 'message', 'popup', 'notification', 'toolbar',
        'window', 'menu', 'panel', 'comment', '弹幕', '弹窗', '聊天', '通知', '窗口', '菜单'
    ]
    score = 0
    for kw in keywords:
        if kw in text:
            score += 1
    faces = res.get('vl', {}).get('faces')
    if isinstance(faces, (int, float)) and faces > 1:
        score += 0.5
    app_name = res.get('vl', {}).get('app_name')
    if isinstance(app_name, str) and any(tag in app_name.lower() for tag in ['zoom', 'teams', 'wechat', 'meeting']):
        score += 0.5
    score += (1.0 - res.get('bestshot_score', 0.0)) * 0.5
    return score


def mark_similar_screenshots(screenshot_items: List[Dict[str, Any]], desc_threshold: float = 0.85):
    """
    对标记为"保留"的截图进行杂乱度去重
    LLM建议删除的截图不参与去重，直接保持删除状态
    """
    if len(screenshot_items) <= 1:
        return
    
    # 只处理 LLM 建议保留的截图（suggestion 不是 'delete' 的）
    keep_items = []
    keep_indices = []
    
    for idx, item in enumerate(screenshot_items):
        # 直接从 item 中读取 LLM 分类结果
        sug_dict = item.get('suggestion', {})
        is_delete = sug_dict.get('delete', False)
        
        # 如果 LLM 建议删除，直接标记为删除，不进入去重流程
        if is_delete:
            # 已经标记过了，跳过
            pass
        else:
            # LLM 建议保留的截图，进入去重流程
            keep_items.append(item)
            keep_indices.append(idx)
    
    # 如果没有需要去重的截图，直接返回
    if len(keep_items) <= 1:
        return
    
    n = len(keep_items)
    parent = list(range(n))

    def find(x):
        while parent[x] != x:
            parent[x] = parent[parent[x]]
            x = parent[x]
        return x

    def union(a, b):
        ra, rb = find(a), find(b)
        if ra != rb:
            parent[rb] = ra

    # 只对保留的截图进行相似度比较
    descriptions = [
        item.get('short_description') or item.get('vl', {}).get('scene') or ''
        for item in keep_items
    ]
    app_names = [item.get('vl', {}).get('app_name') for item in keep_items]

    for i in range(n):
        for j in range(i + 1, n):
            desc_sim = short_desc_similarity(descriptions[i], descriptions[j])
            app_match = app_names[i] and app_names[i] == app_names[j]
            if desc_sim >= desc_threshold or app_match:
                union(i, j)

    groups: Dict[int, List[Dict[str, Any]]] = {}
    for idx in range(n):
        root = find(idx)
        groups.setdefault(root, []).append(keep_items[idx])

    for group in groups.values():
        if len(group) < 2:
            continue
        # 在保留的截图中，选择杂乱度最低的
        best = min(group, key=lambda r: (compute_screenshot_clutter(r), -r.get('bestshot_score', 0)))
        for item in group:
            if item is best:
                continue
            desc_sim = short_desc_similarity(
                item.get('short_description') or '',
                best.get('short_description') or ''
            )
            app_same = (item.get('vl', {}).get('app_name') or '') == (best.get('vl', {}).get('app_name') or '')
            if desc_sim >= 0.98 and app_same:
                sug = item.setdefault('suggestion', {})
                if not sug.get('delete'):
                    sug['delete'] = True
                    sug['reason'] = f"相似截图，保留更干净的版本 {Path(best['path']).name}"
                    item['duplicate_of'] = best['path']


def print_screenshot_tree(results: List[Dict[str, Any]]):
    """显示图片分类树状结构（真实照片 + 非照片四分类）"""
    real_photos_list = []
    non_photo_tree: Dict[str, List[str]] = {
        "临时类": [],
        "参考类": [],
        "收藏类": [],
        "记忆类": []
    }
    
    for res in results:
        is_real_photo = not res.get('likely_screenshot', False)
        filename = Path(res['path']).name
        
        if is_real_photo:
            # 真实照片
            score = res.get('bestshot_score')
            suggestion_text = "删除" if res.get('suggestion', {}).get('delete') else "保留"
            reason = ""
            if res.get('suggestion', {}).get('reason') and res.get('suggestion', {}).get('delete'):
                reason = f" - {res['suggestion']['reason']}"
            
            if score is not None:
                entry = f"{filename} [{suggestion_text}] (质量 {score:.2f}){reason}"
            else:
                entry = f"{filename} [{suggestion_text}]{reason}"
            
            real_photos_list.append(entry)
        else:
            # 非照片类（直接从 res 中读取）
            category_en = res.get('category', 'saved').lower()
            category_cn = CATEGORY_CN.get(category_en, '收藏类')
            
            # 从 suggestion 字典中判断
            sug_dict = res.get('suggestion', {})
            is_delete = sug_dict.get('delete', False)
            suggestion = translate_suggestion('delete' if is_delete else 'keep')
            app_name = res.get('app_name', '')
            
            reason = ""
            if res.get('suggestion', {}).get('reason') and res.get('suggestion', {}).get('delete'):
                reason = f" - {res['suggestion']['reason']}"
            
            # 构建条目：文件名 [建议] (APP名) 原因
            if app_name:
                entry = f"{filename} [{suggestion}] ({app_name}){reason}"
            else:
                entry = f"{filename} [{suggestion}]{reason}"
            
            non_photo_tree[category_cn].append(entry)
    
    # 检查是否有图片
    has_photos = len(real_photos_list) > 0
    has_non_photos = any(len(files) > 0 for files in non_photo_tree.values())
    
    if not has_photos and not has_non_photos:
        return
    
    print("\n📂 图片分类汇总")
    
    # 显示真实照片
    if has_photos:
        print(f"\n📸 【真实照片】({len(real_photos_list)} 张) - 使用 BestShot 去重")
        for entry in sorted(real_photos_list):
            print(f"  - {entry}")
    
    # 显示非照片类（按删除可能性）
    if has_non_photos:
        print(f"\n📱 【非照片类】（按删除可能性排序）")
        
        categories_order = [
            ("临时类", "🗑️", "删除可能性最高"),
            ("参考类", "📋", "删除可能性中高"),
            ("收藏类", "💾", "删除可能性中低"),
            ("记忆类", "❓", "扫描件等重要内容")
        ]
        
        for category, icon, desc in categories_order:
            files = non_photo_tree[category]
            if files:
                print(f"\n{icon} 【{category}】({len(files)} 张) - {desc}")
                for entry in sorted(files):
                    print(f"  - {entry}")


def translate_suggestion(text: str) -> str:
    """将英文建议转换为中文"""
    translations = {"keep": "保留", "delete": "删除"}
    return translations.get(text.lower(), text)


def analyze_similar_files_with_llm(file1: Dict, file2: Dict, similarity: float) -> str:
    """
    调用大模型分析两个相似文件的相同和不同内容
    
    注意：此函数使用原始文本内容（file['content']），而不是embedding向量。
    因为：
    1. Embedding向量是数值向量，无法直接用于LLM的文本分析
    2. LLM需要看到实际文本内容才能理解语义、识别差异、生成详细分析
    
    使用智能截取策略：
    1. 如果文件较短（<15000字符），使用完整内容
    2. 如果文件较长，使用开头+结尾+中间关键部分
    """
    # 智能截取策略：对于长文本，提取开头、结尾和中间部分
    def smart_truncate(content: str, max_chars: int = 15000) -> str:
        """智能截取文本，保留开头、结尾和中间关键部分"""
        if len(content) <= max_chars:
            return content
        
        # 计算各部分长度
        part_size = max_chars // 3  # 每部分约1/3
        
        # 提取开头、中间、结尾
        start_part = content[:part_size]
        middle_start = len(content) // 2 - part_size // 2
        middle_part = content[middle_start:middle_start + part_size]
        end_part = content[-part_size:]
        
        return f"{start_part}\n\n[... 中间部分省略 ...]\n\n{middle_part}\n\n[... 中间部分省略 ...]\n\n{end_part}"
    
    content1 = smart_truncate(file1['content'])
    content2 = smart_truncate(file2['content'])
    
    # 显示实际使用的字符数
    original_len1 = len(file1['content'])
    original_len2 = len(file2['content'])
    used_len1 = len(content1)
    used_len2 = len(content2)
    
    if original_len1 > used_len1 or original_len2 > used_len2:
        print(f"   ℹ️ 文件内容较长，使用智能截取进行分析：")
        print(f"      文件1: {used_len1}/{original_len1} 字符")
        print(f"      文件2: {used_len2}/{original_len2} 字符")
    
    prompt = f"""你是文件内容分析专家。请分析以下两个相似文件的内容，完成以下任务：

    1. 总结两个文件的相同内容（共同点）
    2. 分别总结概述文件1的不同内容（文件1独有的内容）
    3. 分别总结概述文件2的不同内容（文件2独有的内容）

    【文件1】
    文件名：{file1['name']}
    内容长度：{original_len1} 字符
    内容：
    {content1}

    【文件2】
    文件名：{file2['name']}
    内容长度：{original_len2} 字符
    内容：
    {content2}

    【相似度】
    两个文件的余弦相似度为：{similarity:.3f}

    注意：如果文件内容较长，可能只显示了部分内容（开头、中间和结尾部分）。请基于提供的内容进行分析，如果发现内容不完整，请在分析结果中说明。

    请按以下格式输出：
    【相同内容】
    ...

    【文件1的不同内容】
    ...

    【文件2的不同内容】
    ...
    """
    
    try:
        response = chat_client.chat.completions.create(
            model=LLM_MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.0  # 使用0.0确保跨平台结果一致
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        print(f"⚠️ LLM 分析失败: {e}")
        return f"分析失败: {str(e)}"


def process_directory(target_dir: Path = TEST_DIR, log_callback=None):
    """主处理函数，接收目标目录作为参数"""
    def log(msg):
        if log_callback:
            log_callback(msg)
        else:
            print(msg)

    log("🚀 相似文件识别工具启动中...")

    # 二次运行防护：打包/GUI 场景下同一进程可能重复调用本函数。
    # 清理运行态缓存，避免上一次运行的缓存/中间结果影响下一次运行，甚至触发异常。
    try:
        IMAGE_ANALYSIS_CACHE.clear()
        PHOTO_FEATURE_CACHE.clear()
        PHOTO_PHASH_CACHE.clear()
        PHOTO_HIST_CACHE.clear()
    except Exception:
        # 缓存清理失败不应影响主流程
        pass
    
    # 扫描文件
    test_dir_str = str(target_dir)
    log(f"\n📁 正在扫描目录: {test_dir_str}")
    file_infos = scan_files(target_dir)
    
    if len(file_infos) == 0:
        log("❌ 未找到任何文件")
        return []
    
    log(f"✅ 找到 {len(file_infos)} 个文件")
    
    # === 图片深度分析与去重 ===
    image_results = []
    
    # 启用图片分析
    try:
        # 创建临时输出目录用于 analyze_and_report_images
        temp_out = target_dir / "logs"
        temp_out.mkdir(exist_ok=True)
        
        log(f"🖼️ 开始扫描并分析图片 (BestShot 去重)...")
        # 修改 analyze_and_report_images 以返回结构化数据而非仅打印
        # 由于原函数设计为打印，我们这里简单包装一下逻辑
        # 为了不大幅修改原函数签名导致其他地方报错，我们这里直接调用核心逻辑
        
        if image_analyzer:
            images = [f for f in file_infos if Path(f['path']).suffix.lower() in IMAGE_EXTENSIONS]
            if images:
                log(f"   发现 {len(images)} 张图片，正在进行视觉理解与质量评估...")
                analyzed_images = []
                pbar = tqdm(images, desc="分析图片")
                for img in pbar:
                    res = get_or_run_image_analysis(img['path'])
                    
                    # 判断：真实照片 or 非照片类
                    is_real_photo = not res.get('likely_screenshot', False)
                    
                    if is_real_photo:
                        # 真实照片 → 补全特征
                        if not res.get('photo_feature'):
                            res['photo_feature'] = compute_photo_feature(img['path'])
                        pbar.set_postfix_str(f"📸 照片")
                    else:
                        # 非照片类 → LLM 四分类
                        try:
                            pbar.set_postfix_str(f"🤖 分类截图...")
                            llm_result = image_analyzer.screenshot_classify_with_llm(res)
                            res['category'] = llm_result.get('category', 'saved')
                            res['app_name'] = llm_result.get('app_name', '')
                            
                            category_cn = CATEGORY_CN.get(res['category'], '未知')
                            suggestion_text = llm_result.get('suggestion', 'keep')
                            pbar.set_postfix_str(f"✅ {category_cn}")
                            
                            # LLM 建议删除 → 标记
                            if llm_result.get('suggestion') == 'delete':
                                res.setdefault('suggestion', {})
                                res['suggestion']['delete'] = True
                                res['suggestion']['reason'] = f'LLM建议删除（{category_cn}）'
                        except Exception as e:
                            pbar.set_postfix_str(f"⚠️ 分类失败")
                            tqdm.write(f"         ⚠️ LLM分类失败: {e}")
                            import traceback
                            tqdm.write(traceback.format_exc())
                            res['category'] = 'saved'
                            res['app_name'] = ''
                    
                    # 补全 embedding
                    if res.get('embedding'):
                        img['embedding'] = res['embedding']
                    
                    analyzed_images.append(res)
                
                # 执行聚类标记
                real_photos = [r for r in analyzed_images if not r.get('likely_screenshot')]
                non_photos = [r for r in analyzed_images if r.get('likely_screenshot')]
                
                log(f"   📸 真实照片: {len(real_photos)} 张")
                log(f"   📱 截图/非照片: {len(non_photos)} 张")
                
                # 统计截图分类分布
                if non_photos:
                    category_stats = {'temporary': 0, 'reference': 0, 'saved': 0, 'memory': 0}
                    for item in non_photos:
                        cat = item.get('category', 'saved')
                        category_stats[cat] = category_stats.get(cat, 0) + 1
                    log(f"   📊 截图分类统计:")
                    log(f"      🗑️ 临时类: {category_stats['temporary']} 张")
                    log(f"      📋 参考类: {category_stats['reference']} 张")
                    log(f"      💾 收藏类: {category_stats['saved']} 张")
                    log(f"      ❓ 记忆类: {category_stats['memory']} 张")
                
                mark_similar_photos(real_photos, debug=False)
                mark_similar_screenshots(non_photos)
                
                # === 收集照片相似组 ===
                photo_groups = {}
                for item in real_photos:
                    if item.get('duplicate_of'):
                        target = item['duplicate_of']
                        photo_groups.setdefault(target, []).append(item)
                
                for target_path, dup_items in photo_groups.items():
                    # 查找 target 的元信息 (通过 target_path)
                    target_item = next((r for r in real_photos if r['path'] == target_path), None)
                    
                    # 构建文件信息（包含path、suggestion和元信息）
                    files_info = [{
                        "path": target_path, 
                        "name": Path(target_path).name,
                        "size": target_item.get('size', 0) if target_item else 0,
                        "mtime": target_item.get('mtime', 0) if target_item else 0,
                        "suggestion": "保留"
                    }]
                    for dup in dup_items:
                        files_info.append({
                            "path": dup['path'],
                            "name": Path(dup['path']).name,
                            "size": dup.get('size', 0),
                            "mtime": dup.get('mtime', 0),
                            "suggestion": "删除"
                        })
                    image_results.append({
                        "type": "photo_group",
                        "files": files_info,
                        "best_shot": target_path
                    })
                
                # === 收集截图分类结果（按四分类组织） ===
                # 1. 杂乱度去重组（相似截图）
                screenshot_dedup_groups = {}
                for item in non_photos:
                    if item.get('duplicate_of'):
                        target = item['duplicate_of']
                        screenshot_dedup_groups.setdefault(target, []).append(item)
                
                for target_path, dup_items in screenshot_dedup_groups.items():
                    # 查找 target 的元信息
                    target_item = next((r for r in non_photos if r['path'] == target_path), None)
                    
                    files_info = [{
                        "path": target_path, 
                        "name": Path(target_path).name,
                        "size": target_item.get('size', 0) if target_item else 0,
                        "mtime": target_item.get('mtime', 0) if target_item else 0,
                        "suggestion": "保留"
                    }]
                    for dup in dup_items:
                        files_info.append({
                            "path": dup['path'],
                            "name": Path(dup['path']).name,
                            "size": dup.get('size', 0),
                            "mtime": dup.get('mtime', 0),
                            "suggestion": "删除"
                        })
                    image_results.append({
                        "type": "screenshot_dedup_group",
                        "files": files_info,
                        "best_shot": target_path
                    })
                
                # 2. 按四分类收集所有截图（展示给用户决策）
                screenshot_by_category = {
                    'temporary': [],
                    'reference': [],
                    'saved': [],
                    'memory': []
                }
                
                for item in non_photos:
                    # 跳过已经在去重组中的截图
                    if item.get('duplicate_of'):
                        continue
                    
                    sug = item.get('suggestion', {})
                    category = item.get('category', 'saved')
                    
                    # 验证并标准化 category，防止意外值导致 KeyError
                    # 处理可能的旧分类或 LLM 返回的意外值
                    if category not in ['temporary', 'reference', 'saved', 'memory']:
                        category_mapping = {
                            'screenshot': 'temporary',
                            'software': 'temporary',
                            'flowchart': 'reference',
                            'document': 'reference',
                            'photo': 'memory',
                            'other': 'saved'
                        }
                        category = category_mapping.get(category.lower() if isinstance(category, str) else category, 'saved')
                    
                    # 所有截图都收集，不管是否建议删除
                    # 让用户看到完整的分类结果并自己决定
                    is_delete = sug.get('delete', False)
                    screenshot_by_category[category].append({
                        "path": item['path'],
                        "name": Path(item['path']).name,
                        "size": item.get('size', 0),
                        "mtime": item.get('mtime', 0),
                        "suggestion": "删除" if is_delete else "保留",
                        "reason": sug.get('reason', '') if is_delete else f"LLM建议保留"
                    })
                
                # 按分类创建卡片组（只创建有内容的分类）
                category_labels = {
                    'temporary': '🗑️ 临时类截图（删除倾向90%）',
                    'reference': '📋 参考类截图（删除倾向60%）',
                    'saved': '💾 收藏类截图（删除倾向30%）',
                    'memory': '❓ 记忆类截图（删除倾向10%）'
                }
                
                for cat in ['temporary', 'reference', 'saved', 'memory']:
                    if screenshot_by_category[cat]:
                        image_results.append({
                            "type": "screenshot_category",
                            "category": cat,
                            "label": category_labels[cat],
                            "files": screenshot_by_category[cat]
                        })
                
                log(f"✅ 图片分析完成，发现 {len(image_results)} 组结果")
                log(f"   - 相似照片组: {len([r for r in image_results if r['type'] == 'photo_group'])} 组")
                log(f"   - 截图去重组: {len([r for r in image_results if r['type'] == 'screenshot_dedup_group'])} 组")
                log(f"   - 临时类截图: {len(screenshot_by_category['temporary'])} 张")
                log(f"   - 参考类截图: {len(screenshot_by_category['reference'])} 张")
                log(f"   - 收藏类截图: {len(screenshot_by_category['saved'])} 张")
                log(f"   - 记忆类截图: {len(screenshot_by_category['memory'])} 张")
            else:
                log("ℹ️ 未发现图片文件")
    except Exception as e:
        log(f"⚠️ 图片分析过程中出错: {e}")
        import traceback
        log(traceback.format_exc())

    # 仅对非图片文件执行相似度流程
    text_like_files = [
        f for f in file_infos
        if Path(f['path']).suffix.lower() not in IMAGE_EXTENSIONS
    ]

    results = []
    if len(text_like_files) < 2:
        log("ℹ️ 文本文件不足，跳过语义分析")
        similar_pairs = []
    else:
        # 查找相似文件对
        log(f"🔍 正在分析 {len(text_like_files)} 个文本文件的相似度...")
        similar_pairs = find_similar_file_pairs(text_like_files, SIMILARITY_THRESHOLD)

    if not similar_pairs:
        log(f"✅ 未找到相似度 >= {SIMILARITY_THRESHOLD} 的文件对")
    else:
        log(f"✅ 找到 {len(similar_pairs)} 对相似文件")
        log("🤖 正在使用大模型分析相似文件差异...")

        for i, (file1, file2, similarity) in enumerate(similar_pairs, 1):
            log(f"📊 分析第 {i}/{len(similar_pairs)} 对: {file1['name']} <-> {file2['name']}")
            analysis = analyze_similar_files_with_llm(file1, file2, similarity)
            results.append({
                "file1": file1,
                "file2": file2,
                "similarity": similarity,
                "analysis": analysis
            })

        # 保存结果到文件（使用时间戳避免二次运行覆盖/文件占用导致失败）
        try:
            logs_dir = target_dir / "logs"
            logs_dir.mkdir(exist_ok=True)
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_file = logs_dir / f"similar_files_analysis_{ts}.txt"
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write("相似文件分析结果\n")
                f.write("=" * 60 + "\n\n")
                for result in results:
                    f.write(f"文件对: {result['file1']} <-> {result['file2']}\n")
                    f.write(f"相似度: {result['similarity']:.3f}\n")
                    f.write("-" * 60 + "\n")
                    f.write(result['analysis'])
                    f.write("\n\n" + "=" * 60 + "\n\n")
            log(f"✅ 分析报告已保存: {output_file.name}")
        except Exception as e:
            log(f"⚠️ 保存报告失败: {e}")

    # 合并图片结果和文本结果
    return results + image_results


def main():
    """主函数"""
    process_directory()


def diagnose_environment():
    """诊断环境配置，检查所有依赖是否正确安装"""
    print("=" * 60)
    print("🔍 环境诊断")
    print("=" * 60)
    
    # 检查 Python 包
    print("\n📦 Python 包检查:")
    packages = {
        "PyPDF2": PdfReader is not None,
        "pypdfium2": pdfium is not None,
        "PIL/Pillow": Image is not None,
        "numpy": np is not None,
    }
    
    for name, installed in packages.items():
        status = "✅" if installed else "❌"
        print(f"   {status} {name}: {'已安装' if installed else '未安装'}")
    
    # 检查 API 密钥
    print("\n🔑 API 密钥检查:")
    dashscope_key = os.getenv("DASHSCOPE_API_KEY")
    deepseek_key = os.getenv("DEEPSEEK_API_KEY")
    print(f"   {'✅' if dashscope_key else '❌'} DASHSCOPE_API_KEY: {'已设置' if dashscope_key else '未设置'}")
    print(f"   {'✅' if deepseek_key else '❌'} DEEPSEEK_API_KEY: {'已设置' if deepseek_key else '未设置'}")
    
    # OCR 配置
    print("\n🔤 OCR 配置:")
    print(f"   - OCR 模型: {OCR_MODEL} (在线API)")
    print(f"   - PDF 渲染分辨率: {int(PDF_RENDER_SCALE * 72)} DPI")
    print(f"   - PDF OCR 页数限制: {'全部页' if MAX_PDF_OCR_PAGES is None else f'前 {MAX_PDF_OCR_PAGES} 页'}")
    
    # 相似度比较配置
    print("\n🔍 相似度比较配置:")
    print(f"   - 文字层面相似度阈值: {TEXT_SIMILARITY_THRESHOLD} (用于初步筛选)")
    print(f"   - 语义相似度阈值: {SIMILARITY_THRESHOLD} (用于最终判断)")
    
    # Embedding 配置
    print("\n📊 Embedding 配置:")
    print(f"   - 模型: {EMBED_MODEL}")
    print(f"   - 最大 Token 数: {EMBEDDING_MAX_TOKENS}")
    print(f"   - 分块大小: {EMBEDDING_CHUNK_SIZE} 字符")
    print(f"   - 分块重叠: {EMBEDDING_CHUNK_OVERLAP} 字符")
    
    print("\n" + "=" * 60)


if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1 and sys.argv[1] == "--diagnose":
        diagnose_environment()
    else:
        main()

